"""Build the reviewer response letter (docx) from PART 2 of paper/JUTIF_R1_response.md.

PART 1 of that file is the Indonesian working note for the author and is never
included. The output is a standalone document for upload to OJS alongside the
revised manuscript.

    python scripts/build_response_letter.py
"""

import re
from pathlib import Path

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "paper" / "JUTIF_R1_response.md"
OUT = ROOT / "paper" / "JUTIF_R1_response_letter.docx"

FONT = "Times New Roman"
SIZE = Pt(11)


def part2(text: str) -> str:
    start = text.index("# PART 2")
    stop = text.index("## Sisa yang harus")  # author-only checklist, excluded
    return text[start:stop]


def add_runs(par, md: str):
    """Render inline markdown: **bold**, *italic*, `code`."""
    for chunk in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", md):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = par.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            run = par.add_run(chunk[1:-1])
            run.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = par.add_run(chunk[1:-1])
            run.font.name = "Consolas"
        else:
            run = par.add_run(chunk)
        run.font.size = SIZE
        if run.font.name is None:
            run.font.name = FONT


def body(doc, md: str, style=None, space_after=6, indent=None):
    par = doc.add_paragraph(style=style)
    par.paragraph_format.space_after = Pt(space_after)
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent is not None:
        par.paragraph_format.left_indent = Cm(indent)
    add_runs(par, md)
    return par


def heading(doc, md: str, size, space_before=12):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(space_before)
    par.paragraph_format.space_after = Pt(4)
    run = par.add_run(md)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = FONT
    return par


def table(doc, rows):
    tbl = doc.add_table(rows=0, cols=len(rows[0]))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, cells in enumerate(rows):
        row = tbl.add_row().cells
        for cell, md in zip(row, cells):
            cell.text = ""
            par = cell.paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            add_runs(par, md)
            if i == 0:
                for run in par.runs:
                    run.bold = True
    widths = (Cm(1.4), Cm(5.0), Cm(10.6)) if len(rows[0]) == 3 else None
    if widths:
        for row in tbl.rows:
            for cell, w in zip(row.cells, widths):
                cell.width = w
    return tbl


def main():
    md = part2(SRC.read_text(encoding="utf-8"))
    doc = docx.Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = SIZE
    for section in doc.sections:
        section.page_width, section.page_height = Cm(21), Cm(29.7)
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Response to Reviewers")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = FONT

    for line in [
        'Manuscript: "Diagnosing a Register-Pragmatic Blind Spot in Javanese Hate Speech '
        'Detection via LLM-Generated Register-Stratified Stimuli"',
        "Journal: Jurnal Teknik Informatika (JUTIF) — Submission #6393",
        "Decision: Revisions Required (2026-07-24) — Revision 1",
    ]:
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_after = Pt(2)
        add_runs(par, line)

    doc.add_paragraph()

    lines = md.split("\n")
    i = 0
    pending_table = []
    while i < len(lines):
        line = lines[i].rstrip()
        i += 1

        if line.startswith("|"):
            pending_table.append(line)
            if i < len(lines) and lines[i].startswith("|"):
                continue
            rows = [
                [c.strip() for c in r.strip().strip("|").split("|")]
                for r in pending_table
                if not re.match(r"^\|[\s|:-]+\|$", r)
            ]
            table(doc, rows)
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
            pending_table = []
            continue

        if not line.strip() or line.strip() == "---":
            continue
        if line.startswith("# "):
            continue  # the PART 2 banner is replaced by the title block above
        if line.startswith("## "):
            heading(doc, line[3:].strip(), 12)
            continue
        if line.startswith("- "):
            body(doc, line[2:].strip(), style="List Bullet", space_after=2)
            continue
        body(doc, line.strip())

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)

    check = docx.Document(OUT)
    words = sum(len(p.text.split()) for p in check.paragraphs)
    print(f"wrote {OUT.relative_to(ROOT)}")
    print(f"  paragraphs={len(check.paragraphs)} tables={len(check.tables)} words={words}")
    leaked = [p.text for p in check.paragraphs if "**" in p.text or p.text.startswith("|")]
    print(f"  leaked markdown: {len(leaked)}")
    indonesian = [p.text for p in check.paragraphs if "Bapak" in p.text or "sudah" in p.text]
    print(f"  author-only lines leaked: {len(indonesian)}")


if __name__ == "__main__":
    main()
