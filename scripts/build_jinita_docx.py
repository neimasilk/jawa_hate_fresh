"""
Convert paper/draft_jinita_submission.md into the official JINITA Template V4 (eng).docx.

Ground truth for formatting is the template file itself (paper/jinita_guidelines/
JINITA Template V4 eng.docx), inspected directly via its XML:
- Page: A4 portrait, 3cm margins all sides, first-page header/footer differs
  (journal name/ISSN/DOI banner - editorial, left untouched)
- Body text: Times New Roman 10pt, justified, 0.5in first-line indent
- Section headings (1. INTRODUCTION etc.): TNR 10pt bold, ALL CAPS, no indent
- Subsection headings (2.1. Title): TNR 10pt bold, title case
- Front-matter title: TNR 16pt bold, centered
- Author line: TNR 12pt bold, centered
- Affiliation / email: TNR 8pt italic, centered
- ARTICLE INFO / ABSTRACT table: TNR, header labels 9pt bold, abstract body
  10pt, keywords/citation cell 7.5pt
- References: TNR 8pt, hanging indent, italic journal-title run

Do NOT re-run ad hoc; this script is the single source of truth for the
Markdown -> DOCX mapping. Edit and re-run to regenerate the file.
"""
import re
import copy
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

TEMPLATE = "paper/jinita_guidelines/JINITA Template V4 eng.docx"
SRC_MD = "paper/draft_jinita_submission.md"
OUT = "paper/JINITA_Amien_Kanthi_Sijabat_2026.docx"

TNR = "Times New Roman"

TITLE = "Diagnosing a Register-Pragmatic Blind Spot in Javanese Hate Speech Detection"
AUTHORS_LINE = "Mukhlis Amien¹*, Yekti Asmoro Kanthi², Daniel Rudiaman Sijabat³"
AFFIL_LINE = "¹,²,³ Department of Informatics, Universitas Bhinneka Nusantara, Malang, Indonesia"
EMAIL_LINE = "email: ¹amien@ubhinus.ac.id, ²yektiasmoro@ubhinus.ac.id, ³daniel223@ubhinus.ac.id"
IEEE_CITATION = (
    'M. Amien, Y. A. Kanthi, and D. R. Sijabat, "Diagnosing a Register-Pragmatic '
    'Blind Spot in Javanese Hate Speech Detection," Journal of Innovation Information '
    'Technology and Application (JINITA), vol. x, no. x, pp. xx-xx, 2026. '
    '[volume/issue/pages to be assigned by the editor]'
)

# ---------------------------------------------------------------------------
# Inline markdown -> runs
# ---------------------------------------------------------------------------

def set_run_font(run, size=10, bold=False, italic=False, name=TNR):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)


TOKEN_RE = re.compile(r'(\*\*.+?\*\*|\*.+?\*)')


def add_inline_runs(paragraph, text, size=10, base_bold=False, base_italic=False):
    """Parse **bold** / *italic* markdown spans (recursively, so bold-containing-
    italic like '**register is *ngoko* here**' resolves correctly) into runs."""
    parts = TOKEN_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            add_inline_runs(paragraph, part[2:-2], size=size, base_bold=True, base_italic=base_italic)
        elif part.startswith('*') and part.endswith('*') and len(part) >= 2:
            add_inline_runs(paragraph, part[1:-1], size=size, base_bold=base_bold, base_italic=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, bold=base_bold, italic=base_italic)


def strip_md(text):
    return text.replace('**', '').replace('*', '')


# ---------------------------------------------------------------------------
# Load template & wipe placeholder body content
# ---------------------------------------------------------------------------

doc = Document(TEMPLATE)
body = doc.element.body

# --- Front matter: title / authors / affiliation / email ---
paras = doc.paragraphs

def clear_paragraph(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

# index 0 = title
p_title = paras[0]
clear_paragraph(p_title)
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run(TITLE)
set_run_font(r, size=16, bold=True)

# index 2 = authors
p_auth = paras[2]
clear_paragraph(p_auth)
p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_auth.add_run(AUTHORS_LINE)
set_run_font(r, size=12, bold=True)

# index 3 = affiliation (single shared affiliation for all 3 authors)
p_aff1 = paras[3]
clear_paragraph(p_aff1)
p_aff1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_aff1.add_run(AFFIL_LINE)
set_run_font(r, size=8, italic=True)

# index 4 = second affiliation slot -> remove entirely (only one affiliation needed)
p_aff2 = paras[4]
p_aff2._element.getparent().remove(p_aff2._element)

# index 5 = email
p_email = paras[5]
clear_paragraph(p_email)
p_email.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_email.add_run(EMAIL_LINE)
set_run_font(r, size=8, italic=True)

print("Front matter done.")

# ---------------------------------------------------------------------------
# Parse the cleaned markdown to pull out Abstract + Keywords for the table
# ---------------------------------------------------------------------------

md = open(SRC_MD, encoding='utf-8').read()

abs_match = re.search(r'## ABSTRACT\n\n(.+?)\n\n\*\*Keywords:\*\* (.+?)\n', md, re.S)
abstract_text = strip_md(abs_match.group(1).strip().replace('\n', ' '))
keywords_text = abs_match.group(2).strip()
keywords_list = [k.strip() for k in keywords_text.split(';')]

print("Abstract words:", len(abstract_text.split()))
print("Keywords:", keywords_list)

# ---------------------------------------------------------------------------
# Fill ARTICLE INFO / ABSTRACT table (table index 0), drop dummy table (index 1)
# ---------------------------------------------------------------------------

tbl0 = doc.tables[0]
tbl1 = doc.tables[1]
tbl1._element.getparent().remove(tbl1._element)


def set_cell_text(cell, paragraphs_spec):
    """paragraphs_spec: list of (text, size, bold, italic) tuples, one per line."""
    # clear existing paragraphs, keep first one to reuse
    tc = cell._tc
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    for text, size, bold, italic in paragraphs_spec:
        p = tc.add_p if False else None
    # use python-docx add_paragraph on cell (works even with 0 paragraphs left? needs >=1)
    # docx requires at least one paragraph in a cell; add via low-level API
    from docx.oxml.ns import qn as _qn
    for text, size, bold, italic in paragraphs_spec:
        new_p = tc.makeelement(_qn('w:p'), {})
        tc.append(new_p)
    # Now wrap and populate using python-docx Paragraph objects
    from docx.text.paragraph import Paragraph
    for idx, (text, size, bold, italic) in enumerate(paragraphs_spec):
        p_elem = tc.findall(_qn('w:p'))[idx]
        para = Paragraph(p_elem, cell)
        add_inline_runs(para, text, size=size, base_bold=bold, base_italic=italic)


# Article history (editorial placeholder)
set_cell_text(tbl0.cell(1, 0), [
    ("Article history:", 9, True, False),
    ("[To be completed by the editorial office upon review/acceptance]", 7.5, False, True),
])

# Abstract (merged cell spanning rows 1-3)
set_cell_text(tbl0.cell(1, 2), [
    (abstract_text, 10, False, False),
])

# Keywords
kw_paras = [("Keywords:", 9, True, False)] + [(k, 7.5, False, False) for k in keywords_list]
set_cell_text(tbl0.cell(2, 0), kw_paras)

# IEEE citation
set_cell_text(tbl0.cell(3, 0), [
    ("IEEE style in citing this article:", 9, True, False),
    (IEEE_CITATION, 7.5, False, False),
])

print("Front-matter table done.")

# ---------------------------------------------------------------------------
# Wipe old body content (everything between the front-matter table and sectPr)
# ---------------------------------------------------------------------------

body_children = list(body)
tbl0_elem = tbl0._element
sectPr = body.find(qn('w:sectPr'))

start_idx = body_children.index(tbl0_elem) + 1
end_idx = body_children.index(sectPr)

for elem in body_children[start_idx:end_idx]:
    body.remove(elem)

print("Old body content wiped.")

# ---------------------------------------------------------------------------
# Build new body content from the cleaned markdown
# ---------------------------------------------------------------------------

# Reference point: insert new paragraphs right before sectPr
anchor = sectPr


def new_paragraph_before_anchor():
    from docx.oxml.ns import qn as _qn
    p_elem = body.makeelement(_qn('w:p'), {})
    body.insert(list(body).index(anchor), p_elem)
    from docx.text.paragraph import Paragraph
    return Paragraph(p_elem, doc)


def add_heading1(text):
    p = new_paragraph_before_anchor()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    add_inline_runs(p, text.upper(), size=10, base_bold=True)
    return p


def add_heading2(text):
    p = new_paragraph_before_anchor()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    add_inline_runs(p, text, size=10, base_bold=True)
    return p


def add_body_paragraph(text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10):
    p = new_paragraph_before_anchor()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    add_inline_runs(p, text, size=size)
    return p


def add_caption(text, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = new_paragraph_before_anchor()
    p.alignment = align
    p.paragraph_format.space_after = Pt(8)
    add_inline_runs(p, text, size=9)
    return p


def add_image(path):
    p = new_paragraph_before_anchor()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(4.6))
    p.paragraph_format.space_before = Pt(6)


def add_table(header, rows, colwidths=None):
    n_cols = len(header)
    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(p, h, size=9, base_bold=True)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(p, val, size=9)
    # move the table element to just before the anchor paragraph
    tbl_elem = tbl._element
    tbl_elem.getparent().remove(tbl_elem)
    body.insert(list(body).index(anchor), tbl_elem)
    # spacer paragraph after table
    sp = new_paragraph_before_anchor()
    sp.paragraph_format.space_after = Pt(6)
    return tbl


def add_reference(num, text):
    p = new_paragraph_before_anchor()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(4)
    label = f"[{num}]\t"
    r = p.add_run(label)
    set_run_font(r, size=8)
    add_inline_runs(p, text, size=8)
    return p


# ---------------------------------------------------------------------------
# Markdown parser: walk the document line-by-line
# ---------------------------------------------------------------------------

lines = md.split('\n')
i = 0
n = len(lines)
in_abstract_block = False
skip_until_next_h2 = False

FIGURE_RE = re.compile(r'^!\[(.*?)\]\((.+?)\)$')
REF_RE = re.compile(r'^\[(\d+)\]\s+(.*)$')

# Skip the markdown's own H1 title / author / affiliation / email block —
# that content is already placed into the front matter (title paragraphs +
# ARTICLE INFO/ABSTRACT table) above, and must not also flow into the body.
for k, ln in enumerate(lines):
    if ln.startswith('## '):
        i = k
        break

while i < n:
    line = lines[i].rstrip()

    if line.startswith('## '):
        heading_text = line[3:].strip()
        if heading_text == 'ABSTRACT':
            skip_until_next_h2 = True
            i += 1
            continue
        skip_until_next_h2 = False
        if re.match(r'^\d+\.', heading_text):
            add_heading1(heading_text)
        else:
            add_heading1(heading_text)  # REFERENCES / ACKNOWLEDGEMENTS / GEN AI DISCLOSURE STATEMENT
        i += 1
        continue

    if skip_until_next_h2:
        i += 1
        continue

    if line.startswith('### '):
        add_heading2(line[4:].strip())
        i += 1
        continue

    if not line.strip():
        i += 1
        continue

    if line.strip() == '---':
        i += 1
        continue

    m_img = FIGURE_RE.match(line.strip())
    if m_img:
        img_path = "paper/" + m_img.group(2)
        add_image(img_path)
        i += 1
        continue

    # Markdown table: header line starting with '|', next line separator '|---|...'
    if line.strip().startswith('|') and i + 1 < n and re.match(r'^\|[\s:|-]+\|$', lines[i + 1].strip()):
        header = [c.strip() for c in line.strip().strip('|').split('|')]
        j = i + 2
        rows = []
        while j < n and lines[j].strip().startswith('|'):
            row = [c.strip() for c in lines[j].strip().strip('|').split('|')]
            rows.append(row)
            j += 1
        add_table(header, rows)
        i = j
        continue

    m_ref = REF_RE.match(line.strip())
    if m_ref:
        add_reference(m_ref.group(1), m_ref.group(2))
        i += 1
        continue

    # regular paragraph (collect continuation lines until blank)
    para_lines = [line]
    j = i + 1
    while j < n and lines[j].strip() and not lines[j].startswith('#') and not lines[j].strip().startswith('|') \
            and not FIGURE_RE.match(lines[j].strip()) and not REF_RE.match(lines[j].strip()) \
            and lines[j].strip() != '---':
        para_lines.append(lines[j].rstrip())
        j += 1
    full_text = ' '.join(l.strip() for l in para_lines)

    is_caption = full_text.startswith('**Fig.') or full_text.startswith('**Table')
    if is_caption:
        add_caption(full_text, align=WD_ALIGN_PARAGRAPH.CENTER if full_text.startswith('**Fig.') else WD_ALIGN_PARAGRAPH.JUSTIFY)
    else:
        add_body_paragraph(full_text)
    i = j

print("Body content inserted.")

# ---------------------------------------------------------------------------
# Corresponding-author footnote in first-page footer (footer2.xml -> rId13)
# ---------------------------------------------------------------------------

for section in doc.sections:
    try:
        footer = section.first_page_footer
        for p in footer.paragraphs:
            if 'Corresponding Author' in p.text:
                for r in list(p.runs):
                    pass
                run = p.add_run(" Mukhlis Amien, amien@ubhinus.ac.id")
                set_run_font(run, size=8)
    except Exception as e:
        print("Footer edit skipped:", e)

doc.save(OUT)
print("Saved:", OUT)
