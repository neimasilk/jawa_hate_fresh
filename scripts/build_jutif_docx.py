"""
Convert paper/draft_jutif_submission.md into the official JUTIF template
(paper/JUTIF-Template.docx), filled in place with python-docx.

Ground truth for formatting is the template file itself, inspected directly
via its XML (paragraph indices, table shapes, style names, numbering ids,
and the actual run-level overrides used in the template's own filled-in
sample content such as the "Received:/Revised:/..." dates line and the
Author/Affiliation/Email lines, which consistently override the nominal
style sizes to 11pt and, for ABSTRAK TITLE-styled real content, override
the style's default bold=True to bold=False).

Do NOT re-run ad hoc; this script is the single source of truth for the
Markdown -> DOCX mapping for the JUTIF submission. Edit and re-run to
regenerate the file.
"""
import re
import zipfile
from lxml import etree

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

TEMPLATE = "paper/JUTIF-Template.docx"
SRC_MD = "paper/draft_jutif_submission.md"
OUT = "paper/JUTIF_Amien_Kanthi_Sijabat_2026.docx"

TNR = "Times New Roman"
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

TITLE = "Diagnosing a Register-Pragmatic Blind Spot in Javanese Hate Speech Detection"

# ---------------------------------------------------------------------------
# Low-level run / paragraph formatting helpers
# ---------------------------------------------------------------------------


def set_run_font(run, size_pt=None, bold=None, italic=None, superscript=False, name=TNR):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if size_pt is not None:
        run.font.size = Pt(size_pt)
        sz_el = rPr.find(qn('w:sz'))
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            if sz_el is not None:
                sz_el.addnext(szCs)
            else:
                rPr.append(szCs)
        szCs.set(qn('w:val'), str(int(size_pt * 2)))
    if superscript:
        run.font.superscript = True
    return run


TOKEN_RE = re.compile(r'(\*\*.+?\*\*|\*.+?\*)')


def add_inline_runs(paragraph, text, size_pt=None, force_bold=None, base_italic=False):
    """Recursively parse **bold** / *italic* markdown spans into runs.

    force_bold=None means "don't touch bold" (inherit from paragraph style);
    force_bold=True/False explicitly sets it. Markdown **bold** spans force
    True for their sub-run only; everything else keeps the caller's setting.
    """
    parts = TOKEN_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            add_inline_runs(paragraph, part[2:-2], size_pt=size_pt, force_bold=True, base_italic=base_italic)
        elif part.startswith('*') and part.endswith('*') and len(part) >= 2:
            add_inline_runs(paragraph, part[1:-1], size_pt=size_pt, force_bold=force_bold, base_italic=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size_pt=size_pt, bold=force_bold, italic=(True if base_italic else None))


def strip_md(text):
    return text.replace('**', '').replace('*', '')


# ---------------------------------------------------------------------------
# Load template; capture element references up front (before any mutation)
# ---------------------------------------------------------------------------

doc = Document(TEMPLATE)
body = doc.element.body

orig_paras = doc.paragraphs          # captured once; indices below are stable
orig_tables = doc.tables             # [0]=abstract table, [1]=sample (removed), [2]=checklist

tbl_abstract = orig_tables[0]
tbl_checklist = orig_tables[2]

intro_anchor_elem = orig_paras[13]._element        # "INTRODUCTION" (start of removal range)
checklist_anchor_elem = orig_paras[90]._element    # "SUBMISSION CHECK-LIST..." (insertion point)


def clear_paragraph(p):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def add_plain_run(p, text, superscript=False, size_pt=11):
    r = p.add_run(text)
    set_run_font(r, size_pt=size_pt, superscript=superscript)
    return r


# ---------------------------------------------------------------------------
# Front matter
# ---------------------------------------------------------------------------

# para 0: title (JUDUL style; keep style size/bold, just force TNR)
p_title = orig_paras[0]
clear_paragraph(p_title)
r = p_title.add_run(TITLE)
set_run_font(r)

# para 1: authors (AUTHOR style). Template's own filled sample overrides the
# style's 10pt to 11pt on every run; mirrored here. "*1" superscripted as a
# unit per spec (corresponding-author marker travels with the numeral).
p_auth = orig_paras[1]
clear_paragraph(p_auth)
add_plain_run(p_auth, "Mukhlis Amien")
add_plain_run(p_auth, "*1", superscript=True)
add_plain_run(p_auth, ", Yekti Asmoro Kanthi")
add_plain_run(p_auth, "2", superscript=True)
add_plain_run(p_auth, ", Daniel Rudiaman Sijabat")
add_plain_run(p_auth, "3", superscript=True)

# para 2: single shared affiliation (INSTANSI style); delete paras 3,4,5
p_aff = orig_paras[2]
clear_paragraph(p_aff)
add_plain_run(p_aff, "1,2,3", superscript=True)
add_plain_run(p_aff, "Informatics, Universitas Bhinneka Nusantara, Indonesia")

for idx in (3, 4, 5, 6):
    el = orig_paras[idx]._element
    el.getparent().remove(el)
# para 6 was an empty INSTANSI spacer; removed too so exactly one INSTANSI
# paragraph (para 2) remains, per QA requirement.

# para 7: email (EMAIL AUTHOR style)
p_email = orig_paras[7]
clear_paragraph(p_email)
add_plain_run(p_email, "Email: ")
add_plain_run(p_email, "1", superscript=True)
add_plain_run(p_email, "amien@ubhinus.ac.id")

# para 8: received/revised/accepted/published dates (ABSTRAK TITLE style,
# which is bold=True by default -- the template's own filled sample
# explicitly overrides to bold=False; mirrored here).
p_dates = orig_paras[8]
clear_paragraph(p_dates)
r = add_plain_run(p_dates, "Received : -; Revised : -; Accepted : -; Published : -")
set_run_font(r, size_pt=11, bold=False)

# paras 9, 10, 11, 12 kept unchanged (phone number line, cellphone note,
# "Abstract" heading, anchored horizontal-line paragraph).

print("Front matter done.")

# ---------------------------------------------------------------------------
# Abstract + keywords from the markdown
# ---------------------------------------------------------------------------

md = open(SRC_MD, encoding='utf-8').read()

abs_match = re.search(r'## ABSTRACT\n\n(.+?)\n\n\*\*Keywords:\*\* (.+?)\n', md, re.S)
abstract_raw = abs_match.group(1).strip().replace('\n', ' ')
keywords_text = abs_match.group(2).strip()

print("Abstract words:", len(strip_md(abstract_raw).split()))
print("Keywords:", keywords_text)


def rebuild_cell_paragraphs(cell, n):
    tc = cell._tc
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    paras = []
    for _ in range(n):
        p_elem = OxmlElement('w:p')
        tc.append(p_elem)
        paras.append(Paragraph(p_elem, cell))
    return paras


cell00 = tbl_abstract.cell(0, 0)
p_abs, p_blank, p_kw = rebuild_cell_paragraphs(cell00, 3)

p_abs.style = doc.styles['ABSTRAK TITLE']
p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_inline_runs(p_abs, abstract_raw, size_pt=11, force_bold=False)

p_blank.style = doc.styles['ABSTRAK TITLE']

p_kw.style = doc.styles['ABSTRAK TITLE']
p_kw.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p_kw.add_run("Keywords :")
set_run_font(r, size_pt=11, bold=True)
r = p_kw.add_run(" ")
set_run_font(r, size_pt=11, bold=False)
add_inline_runs(p_kw, keywords_text, size_pt=11, force_bold=False)

print("Abstract table front cell done. CC-license cell (1,0) left untouched.")

# ---------------------------------------------------------------------------
# Remove template body content: from INTRODUCTION through (not including)
# the SUBMISSION CHECK-LIST heading. This also removes the sample data
# table (a direct body child inside this range).
# ---------------------------------------------------------------------------

body_children = list(body)
start_idx = body_children.index(intro_anchor_elem)
end_idx = body_children.index(checklist_anchor_elem)
for elem in body_children[start_idx:end_idx]:
    body.remove(elem)

print("Old body content wiped (paras + sample table).")

# ---------------------------------------------------------------------------
# New-content insertion helpers (all inserted immediately before the
# checklist heading, so sequential calls preserve document order)
# ---------------------------------------------------------------------------


def new_paragraph():
    p_elem = OxmlElement('w:p')
    checklist_anchor_elem.addprevious(p_elem)
    return Paragraph(p_elem, doc)


def set_run_szCs(run, val):
    rPr = run._element.get_or_add_rPr()
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(val))
    rPr.append(szCs)


def add_intro_heading(text):
    p = new_paragraph()
    p.style = doc.styles['BODY PARAGRAP']
    pPr = p._p.get_or_add_pPr()
    numPr = pPr.get_or_add_numPr()
    numPr.get_or_add_ilvl().val = 0
    numPr.get_or_add_numId().val = 4
    p.paragraph_format.left_indent = Twips(567)
    p.paragraph_format.first_line_indent = Twips(-578)
    mark_rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), TNR)
    rFonts.set(qn('w:hAnsi'), TNR)
    b = OxmlElement('w:b')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'auto')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '24')
    mark_rPr.append(rFonts)
    mark_rPr.append(b)
    mark_rPr.append(color)
    mark_rPr.append(sz)
    pPr.append(mark_rPr)
    r = p.add_run(text)
    set_run_font(r, size_pt=12, bold=True)
    return p


MAIN_HEADING_STYLES = {
    'METHOD': '3 METODE',
    'RESULT': '4 PEMBAHASAN',
    'DISCUSSION': '4 PEMBAHASAN',
    'CONCLUSION': '5 KESIMPULAN',
}


def add_std_main_heading(text, style_name):
    p = new_paragraph()
    p.style = doc.styles[style_name]
    pPr = p._p.get_or_add_pPr()
    numPr = pPr.get_or_add_numPr()
    numPr.get_or_add_ilvl().val = 0
    numPr.get_or_add_numId().val = 4
    p.paragraph_format.left_indent = Twips(567)
    p.paragraph_format.first_line_indent = Twips(-567)
    mark_rPr = OxmlElement('w:rPr')
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '24')
    mark_rPr.append(szCs)
    pPr.append(mark_rPr)
    r = p.add_run(text)
    set_run_szCs(r, 24)
    return p


TAIL_HEADINGS = {'CONFLICT OF INTEREST', 'GENERATIVE AI DISCLOSURE', 'ACKNOWLEDGEMENT'}


def add_unnumbered_heading(text, style_name):
    p = new_paragraph()
    p.style = doc.styles[style_name]
    pPr = p._p.get_or_add_pPr()
    mark_rPr = OxmlElement('w:rPr')
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '24')
    mark_rPr.append(szCs)
    pPr.append(mark_rPr)
    r = p.add_run(text)
    set_run_szCs(r, 24)
    return p


def add_subsection(text):
    p = new_paragraph()
    p.style = doc.styles['SUB JUDUL']
    pPr = p._p.get_or_add_pPr()
    numPr = pPr.get_or_add_numPr()
    numPr.get_or_add_ilvl().val = 1
    numPr.get_or_add_numId().val = 4
    p.paragraph_format.left_indent = Twips(567)
    p.paragraph_format.first_line_indent = Twips(-567)
    mark_rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'auto')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '22')
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '22')
    mark_rPr.append(color)
    mark_rPr.append(sz)
    mark_rPr.append(szCs)
    pPr.append(mark_rPr)
    add_inline_runs(p, text, size_pt=11)
    return p


def add_body_paragraph(text):
    p = new_paragraph()
    p.style = doc.styles['BODY PARAGRAP']
    p.paragraph_format.first_line_indent = Twips(567)
    add_inline_runs(p, text, size_pt=11)
    return p


def add_caption(raw_line, style_name, before_pt, after_pt):
    text = raw_line.strip().replace('**', '')
    p = new_paragraph()
    p.style = doc.styles[style_name]
    pPr = p._p.get_or_add_pPr()
    numPr = pPr.get_or_add_numPr()
    numPr.get_or_add_ilvl().val = 0
    numPr.get_or_add_numId().val = 0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before_pt)
    p.paragraph_format.space_after = Pt(after_pt)
    add_inline_runs(p, text)
    return p


def add_image(img_path):
    p = new_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Twips(0)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(img_path, width=Inches(6.0))
    return p


EQ1_XML = (
    '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr>'
    '<m:t>α=1−</m:t></m:r>'
    '<m:f><m:fPr><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:fPr>'
    '<m:num><m:sSub><m:sSubPr><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:sSubPr>'
    '<m:e><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>D</m:t></m:r></m:e>'
    '<m:sub><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>o</m:t></m:r></m:sub></m:sSub></m:num>'
    '<m:den><m:sSub><m:sSubPr><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:sSubPr>'
    '<m:e><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>D</m:t></m:r></m:e>'
    '<m:sub><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>e</m:t></m:r></m:sub></m:sSub></m:den></m:f></m:oMath>'
)

EQ2_XML = (
    '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<m:sSub><m:sSubPr><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:sSubPr>'
    '<m:e><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>DR</m:t></m:r></m:e>'
    '<m:sub><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>n,d</m:t></m:r></m:sub></m:sSub>'
    '<m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>=</m:t></m:r>'
    '<m:f><m:fPr><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:fPr>'
    '<m:num><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>1</m:t></m:r></m:num>'
    '<m:den><m:d><m:dPr><m:begChr m:val="|"/><m:endChr m:val="|"/><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:dPr>'
    '<m:e><m:sSub><m:sSubPr><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:sSubPr>'
    '<m:e><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>C</m:t></m:r></m:e>'
    '<m:sub><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>n</m:t></m:r></m:sub></m:sSub></m:e></m:d></m:den></m:f>'
    '<m:nary><m:naryPr><m:chr m:val="∑"/><m:limLoc m:val="undOvr"/><m:supHide m:val="1"/><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:naryPr>'
    '<m:sub><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>c∈</m:t></m:r>'
    '<m:sSub><m:sSubPr><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:sSubPr>'
    '<m:e><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>C</m:t></m:r></m:e>'
    '<m:sub><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>n</m:t></m:r></m:sub></m:sSub></m:sub><m:sup/>'
    '<m:e><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>1</m:t></m:r>'
    '<m:d><m:dPr><m:begChr m:val="["/><m:endChr m:val="]"/><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:dPr>'
    '<m:e><m:sSub><m:sSubPr><m:ctrlPr><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:i/></w:rPr></m:ctrlPr></m:sSubPr>'
    '<m:e><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>v</m:t></m:r></m:e>'
    '<m:sub><m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>d</m:t></m:r></m:sub></m:sSub>'
    '<m:r><w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t>(c)=1</m:t></m:r></m:e></m:d></m:e></m:nary></m:oMath>'
)


def add_equation(xml_str, number):
    p = new_paragraph()
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9298')
    tabs.append(tab)
    pPr.append(tabs)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    oMath = etree.fromstring(xml_str.encode('utf-8'))
    p._p.append(oMath)
    r = p.add_run(f"\t({number})")
    set_run_font(r, size_pt=11)
    return p


def add_reference(num, text):
    p = new_paragraph()
    p.paragraph_format.left_indent = Twips(567)
    p.paragraph_format.first_line_indent = Twips(-567)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.tab_stops.add_tab_stop(Twips(567))
    r = p.add_run(f"[{num}]\t")
    set_run_font(r, size_pt=11)
    add_inline_runs(p, text, size_pt=11)
    return p


def add_page_break():
    p = new_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)
    return p


def mk_border(tag, single):
    b = OxmlElement(f'w:{tag}')
    if single:
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
    else:
        b.set(qn('w:val'), 'nil')
    return b


def set_tc_pr(tc, pct, top_single, bottom_single):
    old = tc.find(qn('w:tcPr'))
    if old is not None:
        tc.remove(old)
    tcPr = OxmlElement('w:tcPr')
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(pct))
    tcW.set(qn('w:type'), 'pct')
    tcPr.append(tcW)
    tcBorders = OxmlElement('w:tcBorders')
    tcBorders.append(mk_border('top', top_single))
    tcBorders.append(mk_border('left', False))
    tcBorders.append(mk_border('bottom', bottom_single))
    tcBorders.append(mk_border('right', False))
    tcPr.append(tcBorders)
    tc.insert(0, tcPr)


def fill_cell(cell, text, bold=False, align_left=False):
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if align_left else WD_ALIGN_PARAGRAPH.CENTER
    add_inline_runs(p, text, size_pt=9, force_bold=(True if bold else None))


def build_table(headers, rows):
    n_cols = len(headers)
    n_rows_total = 1 + len(rows)
    tbl = doc.add_table(rows=n_rows_total, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl._element.getparent().remove(tbl._element)
    checklist_anchor_elem.addprevious(tbl._element)

    pct_total = min(5000, 1800 + 700 * n_cols)
    col_pct = pct_total // n_cols

    tblPr = tbl._element.find(qn('w:tblPr'))
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(pct_total))
    tblW.set(qn('w:type'), 'pct')
    tblPr.insert(0, tblW)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)

    all_rows = [headers] + rows
    for ri, rowdata in enumerate(all_rows):
        top = (ri == 0)
        bottom = (ri == 0 or ri == n_rows_total - 1)
        for ci, val in enumerate(rowdata):
            cell = tbl.cell(ri, ci)
            set_tc_pr(cell._tc, col_pct, top, bottom)
            fill_cell(cell, val, bold=(ri == 0), align_left=(ci == 0 and ri > 0))

    # spacer paragraph after the table so following content doesn't glue to it
    sp = new_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return tbl


NUM_PREFIX_RE = re.compile(r'^\d+(?:\.\d+)*\.\s*')


def strip_heading_num(text):
    return NUM_PREFIX_RE.sub('', text).strip()


# ---------------------------------------------------------------------------
# Markdown body parser
# ---------------------------------------------------------------------------

lines = md.split('\n')
start_idx = next(i for i, l in enumerate(lines) if l.strip() == '## 1. INTRODUCTION')

FIGURE_RE = re.compile(r'^!\[(.*?)\]\((.+?)\)$')
REF_RE = re.compile(r'^\[(\d+)\]\s+(.*)$')
TABLE_SEP_RE = re.compile(r'^\|[\s:|-]+\|$')

i = start_idx
n = len(lines)

subsection_count = 0
table_count = 0
figure_count = 0
eq_count = 0
ref_count = 0

while i < n:
    line = lines[i].rstrip()

    if line.startswith('## '):
        heading_text = strip_heading_num(line[3:].strip())
        if heading_text == 'INTRODUCTION':
            add_intro_heading(heading_text)
        elif heading_text in MAIN_HEADING_STYLES:
            add_std_main_heading(heading_text, MAIN_HEADING_STYLES[heading_text])
        elif heading_text in TAIL_HEADINGS:
            add_unnumbered_heading(heading_text, '6 UCAPAN TERIMA KASIH')
        elif heading_text == 'REFERENCES':
            add_unnumbered_heading(heading_text, '7 DAFTAR PUSTAKA')
        else:
            raise ValueError(f"Unhandled top-level heading: {heading_text!r}")
        i += 1
        continue

    if line.startswith('### '):
        subsection_count += 1
        add_subsection(strip_heading_num(line[4:].strip()))
        i += 1
        continue

    if not line.strip():
        i += 1
        continue

    if line.strip() == '---':
        i += 1
        continue

    m_img = FIGURE_RE.match(line.strip())
    if m_img:
        figure_count += 1
        add_image("paper/" + m_img.group(2))
        i += 1
        continue

    if line.strip().startswith('|') and i + 1 < n and TABLE_SEP_RE.match(lines[i + 1].strip()):
        header = [c.strip() for c in line.strip().strip('|').split('|')]
        j = i + 2
        rows = []
        while j < n and lines[j].strip().startswith('|'):
            rows.append([c.strip() for c in lines[j].strip().strip('|').split('|')])
            j += 1
        table_count += 1
        build_table(header, rows)
        i = j
        continue

    if line.strip().startswith('{{EQUATION_'):
        if 'EQUATION_1' in line:
            eq_count += 1
            add_equation(EQ1_XML, 1)
        elif 'EQUATION_2' in line:
            eq_count += 1
            add_equation(EQ2_XML, 2)
        else:
            raise ValueError(f"Unknown equation placeholder: {line!r}")
        i += 1
        continue

    m_ref = REF_RE.match(line.strip())
    if m_ref:
        ref_count += 1
        add_reference(m_ref.group(1), m_ref.group(2))
        i += 1
        continue

    # regular paragraph / caption (collect continuation lines until a blank
    # line or the start of another structural element)
    para_lines = [line]
    j = i + 1
    while (j < n and lines[j].strip()
           and not lines[j].startswith('#')
           and not lines[j].strip().startswith('|')
           and not FIGURE_RE.match(lines[j].strip())
           and not REF_RE.match(lines[j].strip())
           and lines[j].strip() != '---'
           and not lines[j].strip().startswith('{{EQUATION_')):
        para_lines.append(lines[j].rstrip())
        j += 1
    full_text = ' '.join(l.strip() for l in para_lines)

    if full_text.startswith('**Table'):
        add_caption(full_text, 'JUDUL TABEL', 6, 2)
    elif full_text.startswith('**Figure'):
        add_caption(full_text, 'JUDUL GAMBAR', 0, 12)
    else:
        add_body_paragraph(full_text)
    i = j

add_page_break()

print(f"Body content inserted: {subsection_count} subsections, {table_count} tables, "
      f"{figure_count} figures, {eq_count} equations, {ref_count} references.")

# ---------------------------------------------------------------------------
# Submission checklist: mark 42 item rows with a centered bold "X"
# ---------------------------------------------------------------------------

ITEM_ROWS = (
    [2, 3, 4, 5, 6, 7, 8, 9, 10, 11]
    + [15, 16, 17, 18, 19, 20]
    + [24, 25, 26, 27, 28, 29]
    + [32, 33, 34, 35, 36, 37]
    + [40, 41, 42, 43]
    + [46, 47, 48, 49, 50, 51, 52, 53, 54, 55]
)
assert len(ITEM_ROWS) == 42, len(ITEM_ROWS)

for row_idx in ITEM_ROWS:
    cell = tbl_checklist.rows[row_idx].cells[2]
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('X')
    set_run_font(r, bold=True)

print(f"Checklist: {len(ITEM_ROWS)} rows marked X.")

doc.save(OUT)
print("Saved:", OUT)

# ===========================================================================
# QA
# ===========================================================================
print("\n===== QA =====")

# 1. Reopen with python-docx AND lxml-parse the raw document.xml (well-formed)
qa_doc = Document(OUT)
with zipfile.ZipFile(OUT) as z:
    xml_bytes = z.read('word/document.xml')
qa_tree = etree.fromstring(xml_bytes)
print("1. Reopened OK with python-docx; document.xml well-formed via lxml:", qa_tree.tag.endswith('document'))

# 2. Counts
n_tables = len(qa_doc.tables)
n_images = xml_bytes.count(b'<pic:pic ') if b'<pic:pic ' in xml_bytes else len(qa_tree.findall('.//' + qn('a:blip').replace('a:blip', '')))  # fallback below
# robust inline-image count via blip elements. The template's own kept
# ABSTRAK TITLE paragraph (para 12) contains one pre-existing decorative
# horizontal-line EMF drawing (rId9) that must stay untouched; the 3
# content figures from the markdown are added afterwards with fresh rIds.
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
with zipfile.ZipFile(OUT) as z:
    rels_xml = z.read('word/_rels/document.xml.rels')
rels_tree = etree.fromstring(rels_xml)
target_by_id = {r.get('Id'): r.get('Target') for r in rels_tree}
blip_embeds = [b.get(f'{{{R_NS}}}embed') for b in qa_tree.findall(f'.//{{{A_NS}}}blip')]
content_figs = [rid for rid in blip_embeds if 'fig' not in (target_by_id.get(rid) or '') or True]
n_images_total = len(blip_embeds)
n_images_content = sum(1 for rid in blip_embeds if not target_by_id.get(rid, '').endswith('image1.emf'))
n_omath = len(qa_tree.findall(f'.//{{http://schemas.openxmlformats.org/officeDocument/2006/math}}oMath'))
ref_paras = [p for p in qa_doc.paragraphs if re.match(r'^\[\d+\]', p.text)]
print(f"2. tables={n_tables} (expect 9), content images={n_images_content} (expect 3; "
      f"total blips incl. template's own kept decorative line = {n_images_total}), "
      f"oMath={n_omath} (expect 2), reference paragraphs={len(ref_paras)} (expect 45)")

# 3. Headings
MAIN_HEADINGS = ['INTRODUCTION', 'METHOD', 'RESULT', 'DISCUSSION', 'CONCLUSION']
found_main = {}
for p in qa_doc.paragraphs:
    txt = p.text.strip()
    if txt in MAIN_HEADINGS:
        pPr = p._p.find(qn('w:pPr'))
        numPr = pPr.find(qn('w:numPr')) if pPr is not None else None
        has_num4_ilvl0 = False
        if numPr is not None:
            ilvl = numPr.find(qn('w:ilvl'))
            numId = numPr.find(qn('w:numId'))
            has_num4_ilvl0 = (ilvl is not None and ilvl.get(qn('w:val')) == '0'
                               and numId is not None and numId.get(qn('w:val')) == '4')
        no_digit_prefix = not re.match(r'^\d', txt)
        found_main[txt] = has_num4_ilvl0 and no_digit_prefix
print("3a. Main headings numPr(numId=4,ilvl=0) + no literal digit prefix:", found_main)

subsection_style_count = 0
for p in qa_doc.paragraphs:
    if p.style.name == 'SUB JUDUL':
        pPr = p._p.find(qn('w:pPr'))
        numPr = pPr.find(qn('w:numPr')) if pPr is not None else None
        if numPr is not None:
            ilvl = numPr.find(qn('w:ilvl'))
            if ilvl is not None and ilvl.get(qn('w:val')) == '1':
                subsection_style_count += 1
md_subsection_count = len(re.findall(r'^### ', md, re.M))
print(f"3b. Subsections in md: {md_subsection_count}; SUB JUDUL ilvl=1 paragraphs in docx: {subsection_style_count}")

# 4. No literal markdown / placeholder artifacts
full_text = '\n'.join(p.text for p in qa_doc.paragraphs)
for tbl in qa_doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            full_text += '\n' + cell.text
bad_markers = ['**', '{{', '](figures/', '## ', '§']
bad_found = {m: (m in full_text) for m in bad_markers}
print("4. Literal artifact scan (all should be False):", bad_found)

# 5. Checklist X count
x_count = 0
for row in tbl_checklist.rows:
    if row.cells[2].text.strip() == 'X':
        x_count += 1
print(f"5. Checklist 'X' marks in column 2: {x_count} (expect 42)")

# 6. Abstract cell + CC license
abs_cell_text = qa_doc.tables[0].cell(0, 0).text
cc_cell_text = qa_doc.tables[0].cell(1, 0).text
print("6. Abstract cell contains 'Keywords :':", 'Keywords :' in abs_cell_text)
print("   CC-license cell intact ('Creative Commons' present):", 'Creative Commons' in cc_cell_text or 'open access' in cc_cell_text)

# 7. Body run fonts (sample 30 random body paragraphs)
import random
body_paras = [p for p in qa_doc.paragraphs if p.style.name == 'BODY PARAGRAP' and p.runs]
random.seed(0)
sample = random.sample(body_paras, min(30, len(body_paras)))
bad_font_runs = 0
checked_runs = 0
for p in sample:
    for r in p.runs:
        checked_runs += 1
        if r.font.name != TNR:
            bad_font_runs += 1
print(f"7. Sampled {len(sample)} BODY PARAGRAP paragraphs / {checked_runs} runs; non-TNR runs: {bad_font_runs} (expect 0)")

# 8. Front matter
title_text = qa_doc.paragraphs[0].text
author_para = qa_doc.paragraphs[1]
has_superscript = any(r.font.superscript for r in author_para.runs)
instansi_count = sum(1 for p in qa_doc.paragraphs if p.style.name == 'INSTANSI')
print(f"8. Title correct: {title_text == TITLE!r} -> {title_text == TITLE}")
print(f"   Author line has superscript run(s): {has_superscript}")
print(f"   Remaining INSTANSI paragraphs: {instansi_count} (expect 1)")

print("===== QA done =====")
